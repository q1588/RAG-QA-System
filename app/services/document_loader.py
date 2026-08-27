# -*- coding: utf-8 -*-
"""文档加载：txt/md/pdf/docx -> 纯文本，带异常解析容错（升级需求 #3 后半部分）。

容错策略：
- txt/md：编码逐级回退 utf-8 -> gbk -> latin-1（latin-1 保证任何字节都能解码，绝不 500）；
- pdf：pypdf 逐页 try/except，坏页跳过、好页照常提取；整份提取为空时尝试 pdfplumber
  （若已安装）兜底；仍为空才抛 DocumentParseError；
- docx：python-docx 段落 + 表格文本提取；损坏文件给出包含文件名与原因的友好错误；
- 通用清洗：去除 NUL/控制字节、统一换行、压缩连续空行，避免脏文本污染向量库。
"""
import os
import re

SUPPORTED_EXTS = {".txt", ".md", ".pdf", ".docx"}


class DocumentParseError(ValueError):
    """文档解析失败（损坏/无文本），detail 面向用户友好。"""


def load_document(path: str) -> str:
    """按扩展名读取文档内容，返回清洗后的纯文本。"""
    ext = os.path.splitext(path)[1].lower()
    if ext in (".txt", ".md"):
        return _read_text_file(path)
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    raise ValueError(f"不支持的文件类型：{ext}（仅支持 {'/'.join(sorted(SUPPORTED_EXTS))}）")


def _clean_text(text: str) -> str:
    """通用清洗：NUL/控制字符剔除、换行统一、连续空行压缩。"""
    text = text.replace("\x00", "")
    text = re.sub("[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub("\n{3,}", "\n\n", text)
    return text.strip()


def _read_text_file(path: str) -> str:
    """兼容 utf-8 / gbk / latin-1 编码的文本读取（latin-1 兜底保证不抛异常）。"""
    with open(path, "rb") as f:
        raw = f.read()
    for encoding in ("utf-8", "gbk", "latin-1"):
        try:
            text = raw.decode(encoding)
            return _clean_text(text)
        except (UnicodeDecodeError, LookupError):
            continue
    return _clean_text(raw.decode("utf-8", errors="replace"))


def _extract_pdf(path: str) -> str:
    """PDF 解析：pypdf 逐页容错；空结果时尝试 pdfplumber 兜底。"""
    filename = os.path.basename(path)
    try:
        from pypdf import PdfReader
    except ImportError:
        raise DocumentParseError("缺少 PDF 解析库 pypdf，请 pip install pypdf")

    try:
        reader = PdfReader(path)
    except Exception as e:  # 文件损坏 / 加密等
        raise DocumentParseError(f"PDF 文件无法打开：{filename}（{type(e).__name__}: {e}）")

    pages: list[str] = []
    for idx, page in enumerate(reader.pages, start=1):
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            # 单页解析失败不影响整份文档：跳过该页并记录
            pages.append("")
    text = _clean_text("\n".join(pages))

    if not text:
        # 兜底：pdfplumber 对部分扫描/复杂版式 PDF 提取效果更好
        try:
            import pdfplumber  # noqa: F401
        except ImportError:
            pdfplumber = None
        if pdfplumber is not None:
            try:
                with pdfplumber.open(path) as pdf:
                    text = _clean_text(
                        "\n".join(page.extract_text() or "" for page in pdf.pages)
                    )
            except Exception:
                text = ""

    if not text:
        raise DocumentParseError(f"PDF 未提取到任何文本：{filename}（可能是扫描件或加密文档）")
    return text


def _extract_docx(path: str) -> str:
    """docx 解析：段落 + 表格文本；损坏文件友好报错。"""
    filename = os.path.basename(path)
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise DocumentParseError("缺少 python-docx 库，请 pip install python-docx")

    try:
        doc = DocxDocument(path)
    except Exception as e:
        raise DocumentParseError(f"docx 文件无法打开：{filename}（{type(e).__name__}: {e}）")

    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = _clean_text("\n".join(parts))
    if not text:
        raise DocumentParseError(f"docx 未提取到任何文本：{filename}")
    return text
